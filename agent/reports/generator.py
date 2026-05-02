from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from rich.console import Console

console = Console()


# ── Path resolution ───────────────────────────────────────────────────────────

def _get_project_root() -> Path:
    here = Path(__file__).resolve()
    for parent in here.parents:
        if (parent / "config.yaml").exists():
            return parent
    fallback = Path.home() / "clawstrike"
    if fallback.exists():
        return fallback
    return Path.cwd()


_PROJECT_ROOT   = _get_project_root()
ENGAGEMENTS_DIR = _PROJECT_ROOT / "engagements"
REPORTS_DIR     = _PROJECT_ROOT / "reports"
ENGAGEMENTS_DIR.mkdir(parents=True, exist_ok=True)
REPORTS_DIR.mkdir(parents=True, exist_ok=True)


# ── Severity / CVSS constants ─────────────────────────────────────────────────

CVSS_MAP = {
    "ssh":     ("Medium",   "5.3", "Exposed SSH service — check for weak creds / outdated version"),
    "http":    ("High",     "7.5", "Unencrypted HTTP — data interception risk, directory exposure"),
    "https":   ("Medium",   "5.0", "HTTPS service — check TLS version and certificate validity"),
    "ftp":     ("High",     "7.5", "FTP — often allows anonymous login or transmits creds in cleartext"),
    "mysql":   ("Critical", "9.1", "Database port exposed — check for unauthenticated access"),
    "ms-sql":  ("Critical", "9.1", "MSSQL exposed — check for default credentials"),
    "rdp":     ("High",     "8.1", "RDP exposed — BlueKeep / brute-force risk"),
    "smb":     ("Critical", "9.8", "SMB exposed — EternalBlue / relay attack surface"),
    "smtp":    ("Medium",   "5.3", "SMTP open — check for open relay and user enumeration"),
    "dns":     ("Medium",   "5.3", "DNS exposed — check for zone transfer (AXFR)"),
    "redis":   ("Critical", "9.8", "Redis exposed — often unauthenticated, RCE via config write"),
    "mongodb": ("Critical", "9.8", "MongoDB exposed — often unauthenticated"),
}
DEFAULT_CVSS = ("Low", "3.1", "Service exposed — review necessity and access controls")

SEVERITY_STYLE = {
    "Critical": ("C00000", RGBColor(0xC0, 0x00, 0x00), True),
    "High":     ("FF4000", RGBColor(0xFF, 0x40, 0x00), True),
    "Medium":   ("FFA500", RGBColor(0xFF, 0xA5, 0x00), False),
    "Low":      ("007040", RGBColor(0x00, 0x70, 0x40), True),
}


# ── Low-level docx helpers ────────────────────────────────────────────────────

def _cvss_for_service(service: str) -> tuple:
    s = service.lower()
    for key, val in CVSS_MAP.items():
        if key in s:
            return val
    return DEFAULT_CVSS


def _cvss_severity_label(cvss: float) -> str:
    if cvss >= 9.0:
        return "Critical"
    if cvss >= 7.0:
        return "High"
    if cvss >= 4.0:
        return "Medium"
    return "Low"


def _set_cell_bg(cell, hex_color: str):
    shading = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _style_severity_cell(cell, severity: str):
    if severity not in SEVERITY_STYLE:
        return
    bg_hex, text_rgb, white_text = SEVERITY_STYLE[severity]
    _set_cell_bg(cell, bg_hex)
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if white_text else RGBColor(0x00, 0x00, 0x00)


def _add_poc_step(doc, number: int, label: str, command: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(f"{number}. {label}")
    run.bold = True
    cmd_para = doc.add_paragraph()
    cmd_para.paragraph_format.left_indent = Inches(0.4)
    pPr = cmd_para._p.get_or_add_pPr()
    pPr.append(parse_xml(
        '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:val="clear" w:color="auto" w:fill="F2F2F2"/>'
    ))
    cmd_run = cmd_para.add_run(command)
    cmd_run.font.name = "Courier New"
    cmd_run.font.size = Pt(9)


def _is_md_separator(line: str) -> bool:
    stripped = line.strip()
    if re.match(r"^\|[-| :]+\|$", stripped):
        return True
    if re.match(r"^[-=]{3,}$", stripped):
        return True
    return False


def _strip_inline_md(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*",     r"\1", text)
    text = re.sub(r"__(.+?)__",     r"\1", text)
    text = re.sub(r"_(.+?)_",       r"\1", text)
    text = re.sub(r"`(.+?)`",       r"\1", text)
    return text.strip()


def _render_analysis(doc, text: str):
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if _is_md_separator(line):
            i += 1
            continue
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not _is_md_separator(lines[i]):
                    table_lines.append(lines[i])
                i += 1
            if not table_lines:
                continue
            rows = [[c.strip() for c in tl.strip().strip("|").split("|")] for tl in table_lines]
            if not rows:
                continue
            cols = max(len(r) for r in rows)
            tbl = doc.add_table(rows=len(rows), cols=cols)
            tbl.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx in range(cols):
                    val = row[c_idx] if c_idx < len(row) else ""
                    cell = tbl.cell(r_idx, c_idx)
                    cell.text = _strip_inline_md(val)
                    if r_idx == 0 and cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].bold = True
            doc.add_paragraph()
            continue
        m = re.match(r"^(#{1,3})\s+(.*)", line)
        if m:
            level_str, heading_text = m.group(1), m.group(2)
            p = doc.add_paragraph()
            run = p.add_run(_strip_inline_md(heading_text))
            run.bold = True
            run.font.size = Pt(12 if len(level_str) == 1 else 11)
            i += 1
            continue
        m = re.match(r"^[\-\*]\s+(.*)", line)
        if m:
            doc.add_paragraph(_strip_inline_md(m.group(1)), style="List Bullet")
            i += 1
            continue
        m = re.match(r"^\d+\.\s+(.*)", line)
        if m:
            doc.add_paragraph(_strip_inline_md(m.group(1)), style="List Number")
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        para_lines = []
        while i < len(lines):
            l = lines[i]
            if (not l.strip()
                    or l.strip().startswith("|")
                    or re.match(r"^(#{1,3}|\s*[\-\*]\s|\d+\.)\s", l)
                    or _is_md_separator(l)):
                break
            para_lines.append(_strip_inline_md(l))
            i += 1
        if para_lines:
            doc.add_paragraph(" ".join(para_lines))


def _parse_engagement(md_path: Path) -> dict:
    text = md_path.read_text(encoding="utf-8")
    data = {"target": "", "timestamp": "", "ports": [], "analysis": ""}
    m = re.search(r"\*\*Target:\*\*\s*`([^`]+)`", text)
    if m:
        data["target"] = m.group(1)
    m = re.search(r"\*\*Timestamp:\*\*\s*(.+?)  ", text)
    if m:
        data["timestamp"] = m.group(1).strip()
    for row in re.finditer(r"\|\s*(\d+)\s*\|\s*(\w+)\s*\|\s*(\w*)\s*\|\s*([^|]*)\s*\|", text):
        port, proto, service, version = row.groups()
        data["ports"].append({
            "port":     port,
            "protocol": proto,
            "service":  service,
            "version":  version.strip().strip("—").strip(),
        })
    m = re.search(r"## Agent Analysis\s*\n([\s\S]+)", text)
    if m:
        data["analysis"] = m.group(1).strip()
    return data


# ── New helpers ───────────────────────────────────────────────────────────────

def add_code_block(doc, content: str, label: str):
    """Labeled monospace code block with gray background and border."""
    label_para = doc.add_paragraph()
    label_run = label_para.add_run(f"◆ {label}")
    label_run.bold = True
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    code_para = doc.add_paragraph()
    code_para.paragraph_format.left_indent  = Inches(0.3)
    code_para.paragraph_format.right_indent = Inches(0.3)

    pPr = code_para._p.get_or_add_pPr()
    pPr.append(parse_xml(
        '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:val="clear" w:color="auto" w:fill="F5F5F5"/>'
    ))
    pPr.append(parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top    w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/>'
        '<w:left   w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/>'
        '<w:right  w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/>'
        '</w:pBdr>'
    ))

    lines = content.strip().splitlines()
    if len(lines) > 50:
        lines = lines[:50] + ["... [truncated — see evidence file for full output]"]

    code_run = code_para.add_run("\n".join(lines))
    code_run.font.name = "Courier New"
    code_run.font.size = Pt(8)


def _safe_target(target: str) -> str:
    return target.replace("://", "_").replace("/", "_").replace(":", "_")


def _evidence_dir(target: str) -> Path:
    return ENGAGEMENTS_DIR / _safe_target(target) / "evidence"


def _read_evidence(target: str, filename: str) -> str:
    base = _evidence_dir(target)
    for directory in (base, base / "post_exploit", base / "exploits"):
        p = directory / filename
        if p.exists():
            return p.read_text(encoding="utf-8", errors="replace")
    return f"[evidence file not found: {filename}]"


def _list_evidence_files(target: str) -> list:
    ev_dir = _evidence_dir(target)
    if not ev_dir.exists():
        return []
    return sorted(ev_dir.rglob("*.txt"))


def _evidence_file_for_finding(option) -> str | None:
    svc = option.service.lower()
    if "ftp" in svc:
        return "phase3_ftp_anon.txt"
    if "ssh" in svc:
        return "phase3_ssh_audit.txt"
    if "mysql" in svc or "sql" in svc:
        return "phase3_mysql_creds.txt"
    if "smb" in svc or "samba" in svc:
        return "phase3_smb_enum.txt"
    if "http" in svc:
        return "phase3_web_paths.txt"
    return None


def _recommendation_for_finding(option) -> str:
    svc = option.service.lower()
    ver = option.version
    if "ftp" in svc:
        return f"Upgrade vsftpd / FTP service immediately. Disable anonymous login. Consider replacing FTP with SFTP."
    if "smb" in svc or "samba" in svc:
        return f"Upgrade Samba to 4.x. Apply MS17-010 patch. Disable SMBv1."
    if "mysql" in svc or "sql" in svc:
        return "Set a strong root password. Bind MySQL to localhost only. Restrict remote access."
    if "ssh" in svc:
        return "Enforce key-based authentication. Disable root login. Update to latest OpenSSH."
    if "rdp" in svc:
        return "Apply BlueKeep patch (CVE-2019-0708). Enable NLA. Restrict RDP access via VPN."
    if "http" in svc and "https" not in svc:
        return "Enable HTTPS. Disable HTTP or redirect to HTTPS. Review exposed directories."
    if "redis" in svc:
        return "Bind Redis to localhost. Enable AUTH. Disable CONFIG rewrite command."
    return f"Upgrade {option.service} {ver} to the latest stable version. Review service necessity."


def _build_attack_chain_text(session) -> str:
    lines = ["Complete Attack Chain:\n"]

    if session.open_ports:
        port_count = len(session.open_ports)
        lines.append(f"  [Phase 1] Network scan → {port_count} port(s) discovered")
        lines.append("       ↓")

    if session.services:
        svc_summary = ", ".join(
            f"{info.get('service','?')} {info.get('version','')} on port {port}".strip()
            for port, info in list(session.services.items())[:3]
        )
        if len(session.services) > 3:
            svc_summary += f" (+ {len(session.services) - 3} more)"
        lines.append(f"  [Phase 2] Service ID → {svc_summary}")
        lines.append("       ↓")

    if session.findings:
        lines.append(f"  [Phase 3] Enumeration → {len(session.findings)} finding(s)")
        lines.append("       ↓")

    if session.exploit_options:
        top = session.exploit_options[0]
        lines.append(f"  [Phase 4] CVE match → {top.cve} confirmed (CVSS {top.cvss})")
        lines.append("       ↓")

    if session.shells:
        lines.append(f"  [Phase 5] Exploitation → MSF module executed")
        lines.append("       ↓")
        lines.append(f"  [Phase 6] Shell acquired → ROOT on {session.target}")

    if session.loot:
        lines.append("       ↓")
        lines.append("  [Post-Ex] Full system access confirmed")

    if len(lines) == 1:
        lines.append("  No engagement data recorded in session.")

    return "\n".join(lines)


# ── Delivery package ──────────────────────────────────────────────────────────

def create_delivery_package(session, report_path: Path, output_dir: Path) -> Path:
    package_name = f"ClawStrike_{_safe_target(session.target)}_{datetime.now().strftime('%Y%m%d')}"
    package_dir  = output_dir / package_name
    package_dir.mkdir(parents=True, exist_ok=True)

    shutil.copy2(report_path, package_dir)

    evidence_src = _evidence_dir(session.target)
    evidence_dst = package_dir / "evidence"
    if evidence_src.exists():
        shutil.copytree(str(evidence_src), str(evidence_dst), dirs_exist_ok=True)

    ev_files = list((evidence_dst).rglob("*.txt")) if evidence_dst.exists() else []
    ev_listing = "\n".join(f"  {f.relative_to(evidence_dst)}" for f in sorted(ev_files)) or "  (none)"

    readme = package_dir / "README.txt"
    readme.write_text(
        f"ClawStrike OS — Penetration Test Delivery Package\n"
        f"=================================================\n\n"
        f"Target:   {session.target}\n"
        f"Date:     {datetime.now().strftime('%Y-%m-%d')}\n"
        f"Profile:  {session.profile}\n\n"
        f"Contents:\n"
        f"---------\n"
        f"{report_path.name}\n"
        f"  Main penetration test report.\n"
        f"  Open with Microsoft Word or LibreOffice Writer.\n\n"
        f"evidence/\n"
        f"  Raw tool output from every phase of testing.\n"
        f"  Referenced in the report for full transparency.\n\n"
        f"{ev_listing}\n\n"
        f"Notes:\n"
        f"------\n"
        f"- All findings are documented in the main report\n"
        f"- Evidence files contain unfiltered tool output\n"
        f"- Report cross-references evidence by filename\n"
        f"- Generated by ClawStrike OS\n\n"
        f"Contact: @rootordie\n"
    )

    console.print(
        f"\n[bold green]✅ Delivery package created:[/bold green]\n"
        f"   {package_dir}/\n"
        f"   ├── {report_path.name}\n"
        f"   ├── evidence/\n"
        f"   └── README.txt\n"
    )
    return package_dir


# ── Main generator ────────────────────────────────────────────────────────────

def generate_report(target: str, session=None, output_dir: Path = None) -> str:
    if output_dir is None:
        output_dir = REPORTS_DIR
    output_dir.mkdir(parents=True, exist_ok=True)

    if session is None:
        return _generate_legacy_report(target, output_dir)

    return _generate_session_report(session, output_dir)


def _generate_session_report(session, output_dir: Path) -> str:
    target     = session.target
    safe       = _safe_target(target)
    timestamp  = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    out_path   = output_dir / f"{safe}_{timestamp}.docx"

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.2)
        sec.right_margin  = Inches(1.2)

    # ── Section 1: Cover Page ─────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title_p.add_run("PENETRATION TEST REPORT")
    t.bold = True
    t.font.size = Pt(24)
    t.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub_p.add_run("ClawStrike OS — CONFIDENTIAL")
    s.font.size = Pt(13)
    s.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_paragraph()
    cover_tbl = doc.add_table(rows=5, cols=2)
    cover_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_data = [
        ("Target",          target),
        ("Date",            datetime.utcnow().strftime("%B %d, %Y")),
        ("Profile",         session.profile or "STANDARD"),
        ("Classification",  "CONFIDENTIAL"),
        ("Conducted by",    "ClawStrike OS"),
    ]
    for i, (label, value) in enumerate(cover_data):
        lr = cover_tbl.cell(i, 0).paragraphs[0].add_run(label)
        lr.bold = True
        lr.font.size = Pt(11)
        vr = cover_tbl.cell(i, 1).paragraphs[0].add_run(value)
        vr.font.size = Pt(11)

    doc.add_page_break()

    # ── Section 2: Executive Summary ─────────────────────────────────────────
    doc.add_heading("1. Executive Summary", level=1)

    opts = session.exploit_options
    crits = sum(1 for o in opts if o.cvss >= 9.0)
    highs = sum(1 for o in opts if 7.0 <= o.cvss < 9.0)
    meds  = sum(1 for o in opts if 4.0 <= o.cvss < 7.0)
    lows  = sum(1 for o in opts if o.cvss < 4.0)
    shells_acquired = len(session.shells) > 0
    result_str = "Exploitation was successful — remote shell access was obtained." if shells_acquired \
                 else "No exploitation was conducted during this engagement."

    doc.add_paragraph(
        f"A penetration test was conducted against {target}. "
        f"The assessment identified {len(opts)} exploitable vulnerability(s): "
        f"{crits} Critical, {highs} High, {meds} Medium, and {lows} Low severity. "
        f"{result_str} "
        f"Immediate remediation is recommended for all Critical and High findings."
    )

    # ── Section 3: Scope and Methodology ─────────────────────────────────────
    doc.add_heading("2. Scope and Methodology", level=1)
    scope_data = [
        ("Target",    target),
        ("Scope",     session.scope or "Single host"),
        ("Profile",   session.profile or "STANDARD"),
        ("Phases",    "1 — Discovery, 2 — Service ID, 3 — Enumeration, 4 — CVE Analysis, 5 — Exploitation, 6 — Post-Exploitation"),
        ("Date",      datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")),
    ]
    stbl = doc.add_table(rows=len(scope_data), cols=2)
    stbl.style = "Table Grid"
    for i, (k, v) in enumerate(scope_data):
        kr = stbl.cell(i, 0).paragraphs[0].add_run(k)
        kr.bold = True
        stbl.cell(i, 1).text = v
    doc.add_paragraph()

    # ── Section 4: Discovered Services ───────────────────────────────────────
    doc.add_heading("3. Discovered Services", level=1)

    if session.services:
        headers = ["Port", "Protocol", "Service", "Version", "Risk"]
        svc_tbl = doc.add_table(rows=1, cols=len(headers))
        svc_tbl.style = "Table Grid"
        svc_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = svc_tbl.rows[0]
        for i, h in enumerate(headers):
            _set_cell_bg(hdr.cells[i], "222222")
            r = hdr.cells[i].paragraphs[0].add_run(h)
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for port, info in sorted(session.services.items(), key=lambda x: int(x[0]) if str(x[0]).isdigit() else 0):
            svc     = info.get("service", "unknown")
            ver     = info.get("version", "")
            proto   = info.get("protocol", "tcp")
            severity, cvss, _ = _cvss_for_service(svc)
            row = svc_tbl.add_row()
            for i, val in enumerate([str(port), proto, svc, ver or "—", f"{severity} ({cvss})"]):
                row.cells[i].text = val
                if i == 4:
                    _style_severity_cell(row.cells[i], severity)
        doc.add_paragraph()
    else:
        doc.add_paragraph("No service data recorded in this session.")

    # ── Section 5: Vulnerability Findings ────────────────────────────────────
    doc.add_heading("4. Vulnerability Findings", level=1)

    sorted_opts = sorted(opts, key=lambda o: o.cvss, reverse=True)

    if not sorted_opts:
        doc.add_paragraph("No vulnerability findings were recorded in this session.")
    else:
        for opt in sorted_opts:
            severity = _cvss_severity_label(opt.cvss)
            doc.add_heading(
                f"FINDING #{opt.number} — {opt.title}",
                level=2,
            )
            badge_tbl = doc.add_table(rows=1, cols=3)
            for i, (label, val, sev) in enumerate([
                ("Severity", severity, severity),
                ("CVSS",     str(opt.cvss), severity),
                ("CVE",      opt.cve or "N/A", None),
            ]):
                cell = badge_tbl.cell(0, i)
                cell.text = f"{label}: {val}"
                if sev:
                    _style_severity_cell(cell, sev)
                else:
                    cell.paragraphs[0].runs[0].bold = True
            doc.add_paragraph()

            detail_data = [
                ("Port",        f"{opt.port}/tcp"),
                ("Service",     f"{opt.service} {opt.version}".strip()),
                ("Confidence",  opt.confidence),
                ("Module",      opt.msf_module or "manual"),
                ("Notes",       opt.notes),
            ]
            dtbl = doc.add_table(rows=len(detail_data), cols=2)
            dtbl.style = "Table Grid"
            for i, (k, v) in enumerate(detail_data):
                kr = dtbl.cell(i, 0).paragraphs[0].add_run(k)
                kr.bold = True
                dtbl.cell(i, 1).text = v
            doc.add_paragraph()

            ev_file = _evidence_file_for_finding(opt)
            if ev_file:
                content = _read_evidence(target, ev_file)
                add_code_block(doc, content, f"Evidence — {ev_file}")
                doc.add_paragraph()

            rec_p = doc.add_paragraph()
            rec_run = rec_p.add_run("Recommendation: ")
            rec_run.bold = True
            rec_p.add_run(_recommendation_for_finding(opt))
            doc.add_paragraph()

    # ── Section 6: Exploitation Results ──────────────────────────────────────
    doc.add_heading("5. Exploitation Results", level=1)

    if not session.shells:
        doc.add_paragraph("No exploitation was conducted during this engagement.")
    else:
        for i, shell in enumerate(session.shells, 1):
            doc.add_heading(f"Exploitation Attempt #{i}", level=2)
            matched_opt = next(
                (o for o in opts if o.title == shell.get("exploit")), None
            )
            exp_data = [
                ("Vulnerability", shell.get("exploit", "Unknown")),
                ("CVE",           matched_opt.cve if matched_opt else "N/A"),
                ("CVSS",          str(matched_opt.cvss) if matched_opt else "N/A"),
                ("Target",        f"{shell.get('target', target)}"),
                ("Module",        shell.get("method", "unknown")),
                ("Result",        "SUCCESS — Shell acquired"),
                ("Timestamp",     shell.get("timestamp", "")),
            ]
            etbl = doc.add_table(rows=len(exp_data), cols=2)
            etbl.style = "Table Grid"
            for j, (k, v) in enumerate(exp_data):
                kr = etbl.cell(j, 0).paragraphs[0].add_run(k)
                kr.bold = True
                vc = etbl.cell(j, 1)
                vc.text = v
                if k == "Result":
                    _set_cell_bg(vc, "007040")
                    vc.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    vc.paragraphs[0].runs[0].bold = True
            doc.add_paragraph()

            ev_content = _read_evidence(target, f"exploit_{i}.txt")
            if "[evidence file not found" not in ev_content:
                add_code_block(doc, ev_content, f"Exploitation Output — exploit_{i}.txt")
                doc.add_paragraph()

    # ── Section 7: Post-Exploitation Findings ────────────────────────────────
    doc.add_heading("6. Post-Exploitation Findings", level=1)

    if not session.loot:
        doc.add_paragraph("No post-exploitation was conducted during this engagement.")
    else:
        for item in session.loot:
            add_code_block(doc, item.content, item.label)
            doc.add_paragraph()

    # ── Section 8: Attack Chain ───────────────────────────────────────────────
    doc.add_heading("7. Attack Chain Summary", level=1)
    chain_text = _build_attack_chain_text(session)
    chain_para = doc.add_paragraph()
    chain_run  = chain_para.add_run(chain_text)
    chain_run.font.name = "Courier New"
    chain_run.font.size = Pt(9)
    doc.add_paragraph()

    # ── Section 9: Recommendations ───────────────────────────────────────────
    doc.add_heading("8. Recommendations", level=1)

    if sorted_opts:
        rec_tbl = doc.add_table(rows=1, cols=2)
        rec_tbl.style = "Table Grid"
        for i, h in enumerate(["Finding", "Recommendation"]):
            hc = rec_tbl.cell(0, i)
            _set_cell_bg(hc, "222222")
            hr = hc.paragraphs[0].add_run(h)
            hr.bold = True
            hr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for opt in sorted_opts:
            row = rec_tbl.add_row()
            row.cells[0].text = f"{opt.title} ({opt.cve})"
            row.cells[1].text = _recommendation_for_finding(opt)
        doc.add_paragraph()

    doc.add_heading("General Recommendations", level=2)
    for rec in [
        "Implement a patch management policy — apply security updates within 30 days of release.",
        "Enforce network segmentation to limit lateral movement between zones.",
        "Apply the principle of least privilege to all services and user accounts.",
        "Deploy an IDS/IPS and monitor logs for exploitation indicators.",
        "Conduct quarterly vulnerability assessments and annual penetration tests.",
        "Enable multi-factor authentication on all internet-facing services.",
    ]:
        doc.add_paragraph(rec, style="List Bullet")

    # ── Section 10: Evidence Index ────────────────────────────────────────────
    doc.add_heading("9. Evidence Index", level=1)
    ev_files = _list_evidence_files(target)

    if ev_files:
        doc.add_paragraph(
            f"The following {len(ev_files)} evidence file(s) were collected during the engagement. "
            f"All files are included in the delivery package alongside this report."
        )
        ev_tbl = doc.add_table(rows=1, cols=2)
        ev_tbl.style = "Table Grid"
        for i, h in enumerate(["File", "Description"]):
            hc = ev_tbl.cell(0, i)
            _set_cell_bg(hc, "222222")
            hr = hc.paragraphs[0].add_run(h)
            hr.bold = True
            hr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        ev_desc = {
            "phase1": "Network discovery scan",
            "phase2": "Service identification",
            "phase3_ftp": "FTP anonymous login check",
            "phase3_ssh": "SSH cipher analysis",
            "phase3_mysql": "MySQL credential check",
            "phase3_smb": "SMB enumeration",
            "phase3_web": "Web path discovery",
            "exploit": "Exploitation evidence",
            "loot": "Post-exploitation loot",
            "pivot": "Pivot recon",
        }
        for ev_path in sorted(ev_files):
            rel = str(ev_path.relative_to(_evidence_dir(target)))
            desc = next((v for k, v in ev_desc.items() if k in rel.lower()), "Tool output")
            row = ev_tbl.add_row()
            row.cells[0].text = rel
            row.cells[1].text = desc
        doc.add_paragraph()
    else:
        doc.add_paragraph("No evidence files were collected in this engagement.")

    doc.save(str(out_path))
    console.print(f"[dim green]report saved → {out_path}[/dim green]")

    package_dir = create_delivery_package(session, out_path, output_dir)
    return str(package_dir)


def _generate_legacy_report(target: str, output_dir: Path) -> str:
    """Backward-compatible report from markdown engagement files (no session)."""
    pattern = f"{_safe_target(target)}_*.md"
    files   = sorted(ENGAGEMENTS_DIR.glob(pattern))

    if not files:
        return f"ERROR: no engagement files found for target '{target}' in {ENGAGEMENTS_DIR}"

    all_ports, latest_analysis, latest_timestamp = [], "", ""
    for f in files:
        eng = _parse_engagement(f)
        all_ports.extend(eng["ports"])
        latest_analysis  = eng["analysis"]  or latest_analysis
        latest_timestamp = eng["timestamp"] or latest_timestamp

    seen, unique_ports = set(), []
    for p in all_ports:
        key = (p["port"], p["protocol"])
        if key not in seen:
            seen.add(key)
            unique_ports.append(p)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.2)
        sec.right_margin  = Inches(1.2)

    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title_p.add_run("VULNERABILITY ASSESSMENT REPORT")
    t.bold = True
    t.font.size = Pt(24)
    t.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    doc.add_paragraph()

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub_p.add_run("Penetration Test Report — CONFIDENTIAL")
    s.font.size = Pt(13)
    s.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    for _ in range(3):
        doc.add_paragraph()

    cover_tbl = doc.add_table(rows=4, cols=2)
    cover_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate([
        ("Target",         target),
        ("Date",           datetime.utcnow().strftime("%B %d, %Y")),
        ("Classification", "CONFIDENTIAL"),
        ("Prepared by",    "ClawStrike OS"),
    ]):
        cover_tbl.cell(i, 0).paragraphs[0].add_run(label).bold = True
        cover_tbl.cell(i, 1).text = value
    doc.add_page_break()

    doc.add_heading("1. Executive Summary", level=1)
    severities = [_cvss_for_service(p["service"])[0] for p in unique_ports]
    doc.add_paragraph(
        f"A vulnerability assessment was conducted against {target}. "
        f"{len(unique_ports)} service(s) identified. "
        f"{severities.count('Critical')} Critical, {severities.count('High')} High, "
        f"{severities.count('Medium')} Medium, {severities.count('Low')} Low. "
        "Immediate remediation recommended for Critical and High findings."
    )

    doc.add_heading("2. Scope", level=1)
    s2 = doc.add_table(rows=2, cols=2)
    s2.style = "Table Grid"
    s2.cell(0, 0).text = "Target"
    s2.cell(0, 1).text = target
    s2.cell(1, 0).text = "Scan Date"
    s2.cell(1, 1).text = latest_timestamp or datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
    doc.add_paragraph()

    doc.add_heading("3. Findings", level=1)
    headers = ["Port", "Protocol", "Service", "Version", "Severity", "CVSS", "Description"]
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        _set_cell_bg(hdr_row.cells[i], "222222")
        r = hdr_row.cells[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for p in unique_ports:
        severity, cvss, desc = _cvss_for_service(p["service"])
        row = tbl.add_row()
        for i, val in enumerate([p["port"], p["protocol"], p["service"], p["version"] or "—", severity, cvss, desc]):
            row.cells[i].text = val
            if i in (4, 5):
                _style_severity_cell(row.cells[i], severity)
    doc.add_paragraph()

    doc.add_heading("4. Proof of Concept", level=1)
    for p in unique_ports:
        severity, cvss, desc = _cvss_for_service(p["service"])
        doc.add_heading(f"Port {p['port']}/{p['protocol']} — {p['service'].upper()}", level=2)
        badge = doc.add_table(rows=1, cols=2)
        badge.cell(0, 0).text = f"Severity: {severity}"
        badge.cell(0, 1).text = f"CVSS v3.1: {cvss}"
        _style_severity_cell(badge.cell(0, 0), severity)
        _style_severity_cell(badge.cell(0, 1), severity)
        doc.add_paragraph()
        doc.add_paragraph(desc)
        svc = p["service"].lower()
        if "ssh" in svc:
            _add_poc_step(doc, 1, "Confirm version", f"nmap -sV -p 22 {target}")
            _add_poc_step(doc, 2, "Audit ciphers",   f"ssh-audit {target}")
        elif "http" in svc and "https" not in svc:
            _add_poc_step(doc, 1, "Directory scan",  f"gobuster dir -u http://{target} -w common.txt")
            _add_poc_step(doc, 2, "Vuln scan",        f"nikto -h http://{target}")
        elif "ftp" in svc:
            _add_poc_step(doc, 1, "Anonymous check",  f"nmap --script=ftp-anon -p 21 {target}")
        elif "smb" in svc:
            _add_poc_step(doc, 1, "Enumerate shares", f"smbclient -L //{target} -N")
            _add_poc_step(doc, 2, "EternalBlue check", f"nmap --script=smb-vuln-ms17-010 -p 445 {target}")
        else:
            _add_poc_step(doc, 1, "Version detect",  f"nmap -sV -p {p['port']} {target}")
        doc.add_paragraph()

    doc.add_heading("5. Recommendations", level=1)
    for rec in [
        "Disable non-essential services using firewall rules.",
        "Enforce key-based SSH authentication and disable password login.",
        "Keep all services on the latest patched versions.",
        "Implement network segmentation to limit lateral movement.",
        "Deploy IDS/IPS and monitor for exploitation attempts.",
        "Conduct regular vulnerability assessments and penetration tests.",
    ]:
        doc.add_paragraph(rec, style="List Bullet")

    if latest_analysis:
        doc.add_heading("6. Agent Analysis", level=1)
        _render_analysis(doc, latest_analysis)

    safe = _safe_target(target)
    evidence_dir = ENGAGEMENTS_DIR / safe / "evidence"
    ev_files = sorted(evidence_dir.glob("*.txt")) if evidence_dir.exists() else []
    if ev_files:
        doc.add_heading("7. Evidence", level=1)
        doc.add_paragraph(f"{len(ev_files)} evidence file(s) captured during the engagement.")
        for ev_file in ev_files:
            doc.add_heading(ev_file.name, level=2)
            content = ev_file.read_text(encoding="utf-8", errors="replace")
            if len(content) > 4000:
                content = content[:4000] + f"\n\n[... {len(content) - 4000} bytes truncated ...]"
            run = doc.add_paragraph().add_run(content)
            run.font.name = "Courier New"
            run.font.size = Pt(8)
            doc.add_paragraph()

    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    out_path  = output_dir / f"{safe}_{timestamp}.docx"
    doc.save(str(out_path))
    return str(out_path)
